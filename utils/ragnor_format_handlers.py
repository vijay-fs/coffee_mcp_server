"""
File format handlers for the Ragnor document extraction API.
"""
import os
import io
from typing import List, Dict, Any
from PIL import Image
from pdf2image import convert_from_bytes
import mimetypes
import tempfile

# Document format handlers


class FormatHandler:
    """Base class for document format handlers."""

    @staticmethod
    def identify_format(file_content: bytes, filename: str) -> str:
        """Identify file format from content and extension."""
        mime_type, _ = mimetypes.guess_type(filename)

        if not mime_type:
            # Fallback to binary if mimetype detection fails
            return "application/octet-stream"

        return mime_type

    @staticmethod
    def convert_to_images(file_content: bytes, file_format: str) -> List[Image.Image]:
        """Convert document to list of PIL images for processing."""
        raise NotImplementedError("Subclasses must implement this method")

    @staticmethod
    def extract_metadata(file_content: bytes, file_format: str) -> Dict[str, Any]:
        """Extract metadata from document."""
        return {"format": file_format}


class PDFHandler(FormatHandler):
    """Handler for PDF documents."""

    @staticmethod
    def get_pdf_page_count(file_content: bytes) -> int:
        """Get the total number of pages in a PDF without loading them all."""
        try:
            # Import here to avoid circular imports
            from PyPDF2 import PdfReader
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            total_pages = len(pdf_reader.pages)
            print(f"PDF has {total_pages} pages in total")
            return total_pages
        except Exception as e:
            print(f"Error reading PDF page count: {str(e)}")
            return 0
            
    @staticmethod
    def convert_to_images(file_content: bytes, file_format: str) -> List[Image.Image]:
        """Convert PDF to list of PIL images.
        
        For large PDFs, this method will return a list with only the first page image,
        and the caller should use convert_page_to_image to get additional pages.
        """
        try:
            print(f"Converting PDF to images with high resolution (DPI=300)")
            
            # Get total number of pages without loading them all at once
            total_pages = PDFHandler.get_pdf_page_count(file_content)
            
            # For large PDFs, just convert the first page and return it
            # The caller should use convert_page_to_image for additional pages
            if total_pages > 10:  # Only convert first page for large PDFs
                print(f"Large PDF detected with {total_pages} pages - returning only first page")
                page_images = convert_from_bytes(
                    file_content,
                    dpi=300,  # Higher DPI for better OCR results
                    fmt='png',
                    thread_count=2,
                    use_cropbox=True,
                    strict=False,
                    transparent=False,
                    grayscale=True,
                    first_page=1,
                    last_page=1
                )
                
                # Debug: Save first page image for debugging only if debug path is configured and not empty
                if page_images and len(page_images) > 0:
                    debug_path_env = os.environ.get('RAGNOR_DEBUG_IMAGES_PATH')
                    if debug_path_env and debug_path_env.strip():
                        # Use exactly the path provided in the environment variable
                        debug_dir = debug_path_env
                        os.makedirs(debug_dir, exist_ok=True)
                        debug_path = os.path.join(debug_dir, "pdf_conversion_page_1.png")
                        page_images[0].save(debug_path)
                        print(f"Saved PDF conversion image to {debug_path} - Size: {page_images[0].size}")
                    
                    print(f"Successfully converted PDF page 1/{total_pages}")
                    return page_images
                else:
                    raise ValueError("Failed to convert first page of PDF")
            
            # For smaller PDFs, convert all pages at once (original behavior)
            images = convert_from_bytes(
                file_content,
                dpi=300,  # Higher DPI for better OCR results
                fmt='png',
                thread_count=2,
                use_cropbox=True,
                strict=False,
                transparent=False,
                grayscale=True  # Keep color for better OCR in some cases
            )

            # Verify images were created successfully
            if not images:
                raise ValueError("No images were extracted from the PDF")

            # Debug: Save the first image for debugging only if debug path is configured and not empty
            if images:
                debug_path_env = os.environ.get('RAGNOR_DEBUG_IMAGES_PATH')
                if debug_path_env and debug_path_env.strip():
                    # Use exactly the path provided in the environment variable
                    debug_dir = debug_path_env
                    os.makedirs(debug_dir, exist_ok=True)
                    # Save first 2 pages for debugging
                    for i, img in enumerate(images[:2]):
                        debug_path = os.path.join(debug_dir, f"pdf_conversion_page_{i+1}.png")
                        img.save(debug_path)
                        print(f"Saved PDF conversion image to {debug_path} - Size: {img.size}")
            
            print(f"Successfully converted PDF to {len(images)} images in total")
            return images

        except Exception as e:
            print(f"Error converting PDF to images: {str(e)}")
            raise ValueError(f"Failed to convert PDF to images: {str(e)}")
    
    @staticmethod
    def convert_page_to_image(file_content: bytes, page_num: int, total_pages: int) -> Image.Image:
        """Convert a specific PDF page to an image."""
        try:
            print(f"Converting PDF page {page_num}/{total_pages} to image")
            page_images = convert_from_bytes(
                file_content,
                dpi=300,  # Higher DPI for better OCR results
                fmt='png',
                thread_count=2,
                use_cropbox=True,
                strict=False,
                transparent=False,
                grayscale=True,
                first_page=page_num,
                last_page=page_num
            )
            
            if not page_images or len(page_images) == 0:
                raise ValueError(f"Failed to convert page {page_num}")
                
            # Debug: Save image for debugging only if debug path is configured and not empty
            if page_num <= 2:  # Only save first 2 pages for debugging
                debug_path_env = os.environ.get('RAGNOR_DEBUG_IMAGES_PATH')
                if debug_path_env and debug_path_env.strip():
                    # Use exactly the path provided in the environment variable
                    debug_dir = debug_path_env
                    os.makedirs(debug_dir, exist_ok=True)
                    debug_path = os.path.join(debug_dir, f"pdf_conversion_page_{page_num}.png")
                    page_images[0].save(debug_path)
                    print(f"Saved PDF conversion image to {debug_path} - Size: {page_images[0].size}")
            
            print(f"Successfully converted PDF page {page_num}/{total_pages} to image")
            return page_images[0]
        except Exception as e:
            print(f"Error converting PDF page {page_num}: {str(e)}")
            raise ValueError(f"Failed to convert PDF page {page_num}: {str(e)}")

    @staticmethod
    def extract_metadata(file_content: bytes, file_format: str) -> Dict[str, Any]:
        """Extract metadata from PDF."""
        from PyPDF2 import PdfReader

        metadata = {"format": "pdf"}
        try:
            pdf = PdfReader(io.BytesIO(file_content))
            metadata["pages"] = len(pdf.pages)

            if pdf.metadata:
                pdf_info = pdf.metadata
                # Extract standard PDF metadata fields
                for key in pdf_info:
                    clean_key = key.strip('/')
                    metadata[clean_key] = str(pdf_info[key])

            return metadata
        except Exception as e:
            print(f"Error extracting PDF metadata: {str(e)}")
            return metadata


class DocxHandler(FormatHandler):
    """Handler for DOCX documents."""

    @staticmethod
    def convert_to_images(file_content: bytes, file_format: str) -> List[Image.Image]:
        """Convert DOCX to list of PIL images."""
        # For DOCX, we need to write to a temp file first
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_path = temp_file.name
            temp_file.write(file_content)

        images = []
        try:
            # Use LibreOffice to convert DOCX to PDF
            pdf_path = temp_path.replace('.docx', '.pdf')
            os.system(
                f'soffice --headless --convert-to pdf --outdir {os.path.dirname(temp_path)} {temp_path}')

            # Check if PDF was created
            if not os.path.exists(pdf_path):
                raise ValueError("Failed to convert DOCX to PDF")

            # Now convert the PDF to images
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
                images = PDFHandler.convert_to_images(
                    pdf_content, 'application/pdf')

            # Clean up temporary files
            os.remove(pdf_path)

        except Exception as e:
            raise ValueError(f"Failed to convert DOCX to images: {str(e)}")
        finally:
            os.remove(temp_path)

        return images

    @staticmethod
    def extract_metadata(file_content: bytes, file_format: str) -> Dict[str, Any]:
        """Extract metadata from DOCX."""
        import docx

        metadata = {"format": "docx"}
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_path = temp_file.name
                temp_file.write(file_content)

            doc = docx.Document(temp_path)

            # Extract document properties
            metadata["pages"] = len(doc.sections)
            metadata["paragraphs"] = len(doc.paragraphs)
            metadata["tables"] = len(doc.tables)

            # Extract core properties
            core_properties = doc.core_properties
            if core_properties:
                metadata["title"] = core_properties.title
                metadata["author"] = core_properties.author
                metadata["created"] = str(
                    core_properties.created) if core_properties.created else None
                metadata["modified"] = str(
                    core_properties.modified) if core_properties.modified else None
                metadata["last_modified_by"] = core_properties.last_modified_by
                metadata["revision"] = core_properties.revision

            # Clean up temporary file
            os.remove(temp_path)

            return metadata
        except Exception as e:
            print(f"Error extracting DOCX metadata: {str(e)}")
            return metadata


class XlsxHandler(FormatHandler):
    """Handler for XLSX documents."""

    @staticmethod
    def convert_to_images(file_content: bytes, file_format: str) -> List[Image.Image]:
        """Convert XLSX to list of PIL images."""
        # For XLSX, we need to write to a temp file first
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as temp_file:
            temp_path = temp_file.name
            temp_file.write(file_content)

        images = []
        try:
            # Use LibreOffice to convert XLSX to PDF
            pdf_path = temp_path.replace('.xlsx', '.pdf')
            os.system(
                f'soffice --headless --convert-to pdf --outdir {os.path.dirname(temp_path)} {temp_path}')

            # Check if PDF was created
            if not os.path.exists(pdf_path):
                raise ValueError("Failed to convert XLSX to PDF")

            # Now convert the PDF to images
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
                images = PDFHandler.convert_to_images(
                    pdf_content, 'application/pdf')

            # Clean up temporary files
            os.remove(pdf_path)

        except Exception as e:
            raise ValueError(f"Failed to convert XLSX to images: {str(e)}")
        finally:
            os.remove(temp_path)

        return images

    @staticmethod
    def extract_metadata(file_content: bytes, file_format: str) -> Dict[str, Any]:
        """Extract metadata from XLSX."""
        import openpyxl

        metadata = {"format": "xlsx"}
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as temp_file:
                temp_path = temp_file.name
                temp_file.write(file_content)

            wb = openpyxl.load_workbook(temp_path, read_only=True)

            # Extract workbook properties
            metadata["sheets"] = len(wb.sheetnames)
            metadata["sheet_names"] = wb.sheetnames

            # Extract document properties
            if wb.properties:
                props = wb.properties
                metadata["title"] = props.title
                metadata["creator"] = props.creator
                metadata["created"] = str(
                    props.created) if props.created else None
                metadata["modified"] = str(
                    props.modified) if props.modified else None
                metadata["lastModifiedBy"] = props.lastModifiedBy
                metadata["revision"] = props.revision

            # Clean up temporary file
            os.remove(temp_path)

            return metadata
        except Exception as e:
            print(f"Error extracting XLSX metadata: {str(e)}")
            return metadata


class PptHandler(FormatHandler):
    """Handler for PPT/PPTX documents."""

    @staticmethod
    def convert_to_images(file_content: bytes, file_format: str) -> List[Image.Image]:
        """Convert PPTX to list of PIL images."""
        # For PPTX, we need to write to a temp file first
        suffix = '.pptx' if 'pptx' in file_format else '.ppt'
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
            temp_path = temp_file.name
            temp_file.write(file_content)

        images = []
        try:
            # Use LibreOffice to convert PPTX to PDF
            pdf_path = temp_path.replace(suffix, '.pdf')
            os.system(
                f'soffice --headless --convert-to pdf --outdir {os.path.dirname(temp_path)} {temp_path}')

            # Check if PDF was created
            if not os.path.exists(pdf_path):
                raise ValueError(f"Failed to convert {suffix} to PDF")

            # Now convert the PDF to images
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
                images = PDFHandler.convert_to_images(
                    pdf_content, 'application/pdf')

            # Clean up temporary files
            os.remove(pdf_path)

        except Exception as e:
            raise ValueError(f"Failed to convert {suffix} to images: {str(e)}")
        finally:
            os.remove(temp_path)

        return images

    @staticmethod
    def extract_metadata(file_content: bytes, file_format: str) -> Dict[str, Any]:
        """Extract metadata from PPTX."""
        import pptx

        metadata = {"format": "pptx"}
        try:
            # Create temporary file
            suffix = '.pptx' if 'pptx' in file_format else '.ppt'
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
                temp_path = temp_file.name
                temp_file.write(file_content)

            presentation = pptx.Presentation(temp_path)

            # Extract document properties
            metadata["slides"] = len(presentation.slides)

            # Extract core properties
            if presentation.core_properties:
                props = presentation.core_properties
                metadata["title"] = props.title
                metadata["author"] = props.author
                metadata["created"] = str(
                    props.created) if props.created else None
                metadata["modified"] = str(
                    props.modified) if props.modified else None
                metadata["last_modified_by"] = props.last_modified_by
                metadata["revision"] = props.revision

            # Clean up temporary file
            os.remove(temp_path)

            return metadata
        except Exception as e:
            print(f"Error extracting PPTX metadata: {str(e)}")
            return metadata


class ImageHandler(FormatHandler):
    """Handler for image documents (JPG, PNG, etc)."""

    @staticmethod
    def convert_to_images(file_content: bytes, file_format: str) -> List[Image.Image]:
        """Convert image to list of PIL images."""
        try:
            # Open the image
            image = Image.open(io.BytesIO(file_content))
            return [image]
        except Exception as e:
            raise ValueError(f"Failed to process image: {str(e)}")

    @staticmethod
    def extract_metadata(file_content: bytes, file_format: str) -> Dict[str, Any]:
        """Extract metadata from image."""
        metadata = {"format": file_format.split('/')[-1]}
        try:
            image = Image.open(io.BytesIO(file_content))
            metadata["width"] = image.width
            metadata["height"] = image.height
            metadata["mode"] = image.mode
            metadata["format"] = image.format

            # Extract EXIF data if available
            if hasattr(image, "_getexif") and image._getexif():
                exif = image._getexif()
                metadata["exif"] = {}
                for tag, value in exif.items():
                    if tag in exif_tags:
                        metadata["exif"][exif_tags[tag]] = str(value)

            return metadata
        except Exception as e:
            print(f"Error extracting image metadata: {str(e)}")
            return metadata


# EXIF tags mapping
exif_tags = {
    271: 'Make',
    272: 'Model',
    306: 'DateTime',
    36867: 'DateTimeOriginal',
    37522: 'DateTimeDigitized',
    34850: 'ExposureProgram',
    34855: 'ISOSpeedRatings',
    37385: 'Flash',
    37386: 'FocalLength',
}

# Format handler registry
FORMAT_HANDLERS = {
    'application/pdf': PDFHandler,
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocxHandler,
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': XlsxHandler,
    'application/vnd.openxmlformats-officedocument.presentationml.presentation': PptHandler,
    'image/jpeg': ImageHandler,
    'image/png': ImageHandler,
    'image/tiff': ImageHandler,
}


def get_handler_for_format(file_format: str) -> FormatHandler:
    """Get the appropriate handler for a file format."""
    handler = FORMAT_HANDLERS.get(file_format)
    if not handler:
        # Try to find a handler for the general type
        general_type = file_format.split('/')[0]
        for fmt, fmt_handler in FORMAT_HANDLERS.items():
            if fmt.startswith(general_type):
                return fmt_handler

    if not handler:
        raise ValueError(f"Unsupported file format: {file_format}")

    return handler
